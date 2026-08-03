import fitz
import os
import re
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from src.config import Config


class SemanticMarkdownSplitter:
    """语义切块：按 Markdown 标题组织语义单元，超长章节按段落/句子兜底。

    - 每个 chunk 带上标题路径作为上下文前缀（例如 "第1章 计算机系统概述 / 1.1.1 操作系统的概念"）
    - 过滤文件头部的 OCR/生成元数据（YAML front matter、AIGC 块）
    - 只保留有实际内容的块，丢弃纯图注/噪声
    """

    def __init__(self, max_chars=Config.CHUNK_SIZE, overlap=Config.CHUNK_OVERLAP):
        self.max_chars = max_chars
        self.overlap = overlap
        self.min_chars = 25

    def split_text(self, text, metadata=None):
        text = self._strip_front_matter(text)
        if not text.strip():
            return []

        sections = []          # [(heading_path, body)]
        heading_parts = []     # [(level, title)]
        body = []

        def flush():
            if not body:
                return
            content = "\n".join(body).strip()
            body.clear()
            if not content:
                return
            path = " / ".join(t for _, t in heading_parts)
            sections.append((path, content))

        for ln in text.splitlines():
            m = re.match(r"^(#{1,4})\s+(.*)$", ln)
            if m:
                flush()
                level = len(m.group(1))
                title = m.group(2).strip()
                heading_parts = [(lv, t) for lv, t in heading_parts if lv < level]
                heading_parts.append((level, title))
            else:
                body.append(ln)
        flush()

        docs = []
        for path, content in sections:
            if len(content) < self.min_chars and not self._has_cjk(content):
                continue
            for piece in self._chunk_content(content):
                full = f"{path}\n\n{piece}" if path else piece
                meta = dict(metadata or {})
                meta["heading"] = path
                meta["chunk_mode"] = "semantic"
                docs.append(Document(page_content=full, metadata=meta))
        return docs

    # ---- 内部工具 ----

    def _strip_front_matter(self, text):
        # 去掉文件开头的 YAML front matter（--- ... ---）
        m = re.match(r"^---\s*\n.*?^---\s*\n?", text, re.S | re.M)
        if m:
            text = text[m.end():]
        # 去掉 OCR/生成器元数据行
        lines = [
            ln for ln in text.splitlines()
            if not re.match(r"^(AIGC:|Label:|ContentProducer:|ProduceID:)", ln.strip())
        ]
        return "\n".join(lines)

    def _chunk_content(self, content):
        """长内容切块：先按段落，再按句子，最后按目标长度合并（带重叠）。"""
        if len(content) <= self.max_chars:
            return [content]

        paras = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
        units = []
        for p in paras:
            if len(p) <= self.max_chars:
                units.append(p)
            else:
                units.extend(self._split_by_sentence(p))

        chunks = []
        cur = ""
        for u in units:
            if not cur:
                cur = u
            elif len(cur) + 1 + len(u) <= self.max_chars:
                cur = f"{cur}\n{u}"
            else:
                chunks.append(cur)
                tail = cur[-self.overlap:] if self.overlap > 0 else ""
                cur = (f"{tail}\n" if tail else "") + u
        if cur:
            chunks.append(cur)
        return chunks

    def _split_by_sentence(self, text):
        parts = re.split(r"(?<=[。！？；])\s*", text)
        sents = [s.strip() for s in parts if s.strip()]
        units = []
        cur = ""
        for s in sents:
            if cur and len(cur) + len(s) > self.max_chars:
                units.append(cur)
                cur = s
            else:
                cur += s
        if cur:
            units.append(cur)
        return units

    @staticmethod
    def _has_cjk(text):
        return any("\u4e00" <= ch <= "\u9fff" for ch in text)


class DocumentParser:
    def __init__(self):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", " "]
        )
        self.chunk_mode = Config.CHUNK_MODE
        self.semantic_splitter = (
            SemanticMarkdownSplitter() if self.chunk_mode == "semantic" else None
        )
    # ===== 统一入口：根据文件后缀自动分流 =====
    def parse(self, file):
        if isinstance(file, str):
            ext = os.path.splitext(file)[1].lower()
        else:
            ext = os.path.splitext(file.name)[1].lower()

        if ext == ".pdf":
            return self.parse_pdf(file)
        elif ext in (".txt", ".md"):
            return self.parse_txt(file)
        elif ext == ".docx":
            return self.parse_docx(file)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")

    # ===== PDF解析（支持路径 + 上传）=====
    def parse_pdf(self, file):
        if isinstance(file, str):
            doc = fitz.open(file)
            file_name = os.path.basename(file)
        else:
            doc = fitz.open(stream=file.read(), filetype="pdf")
            file_name = file.name

        # PDF 类型检测：扫描件还是文本型
        sample_pages = min(10, doc.page_count)
        text_pages = 0
        for i in range(sample_pages):
            if doc[i].get_text().strip():
                text_pages += 1

        if text_pages == 0:
            raise ValueError("当前仅支持文本型PDF，暂不支持扫描图片型PDF")

        documents = []

        for page_num, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file_name,
                            "page": page_num + 1,
                            "type": "pdf"
                        }
                    )
                )

        return self.splitter.split_documents(documents)

    # ===== TXT/Markdown 解析（支持路径字符串 + 上传文件对象） =====
    def parse_txt(self, file):
        if isinstance(file, str):
            with open(file, "r", encoding="utf-8") as f:
                text = f.read()
            file_name = os.path.basename(file)
        else:
            text = file.read().decode("utf-8")
            file_name = file.name

        if self.semantic_splitter is not None:
            return self.semantic_splitter.split_text(
                text, {"source": file_name, "page": 1, "type": "md"}
            )

        return self.splitter.create_documents(
            [text],
            metadatas=[{"source": file_name, "page": 1, "type": "md"}]
        )

    # ===== DOCX 解析（段落 + 表格） =====
    def parse_docx(self, file):
        docx = DocxDocument(file)
        file_name = file.name

        parts = []

        # 1. 提取所有段落文本
        paragraph_texts = []
        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                paragraph_texts.append(text)

        if paragraph_texts:
            parts.append("\n".join(paragraph_texts))

        # 2. 提取所有表格内容
        for table in docx.tables:
            table_lines = []
            if table.rows:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                for row in table.rows[1:]:
                    cell_texts = [cell.text.strip() for cell in row.cells]
                    if len(header_cells) == len(cell_texts):
                        row_text = " | ".join(
                            f"{h}: {v}" for h, v in zip(header_cells, cell_texts) if v
                        )
                    else:
                        row_text = " | ".join(filter(None, cell_texts))
                    if row_text.strip():
                        table_lines.append(row_text)

            if table_lines:
                parts.append("\n".join(table_lines))

        if not parts:
            return []

        combined_text = "\n\n".join(parts)

        return self.splitter.create_documents(
            [combined_text],
            metadatas=[{"source": file_name, "page": "N/A", "type": "docx"}]
        )
