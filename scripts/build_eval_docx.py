"""生成 408 RAG 评测总报告 DOCX。

设计预设：compact_reference_guide（Calibri 11pt / 正文 after 6pt / 行距 1.25；
H1 16pt #2E74B5，H2 13pt #2E74B5，H3 12pt #1F4D78；表格 9360 DXA、表头 #E8EEF5）。
中文正文 East Asia 字体使用微软雅黑（预设的命名覆盖）。

用法（需 python-docx）：
    python scripts/build_eval_docx.py [输出路径]
数据源：results/ 下各轮评测 JSON（直接读取，避免转写错误）。
"""

import json
import copy
import os
import statistics
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, "408RAG评测报告.docx")

# ==================== 预设 token（compact_reference_guide）====================
BODY_FONT = "Calibri"
CJK_FONT = "微软雅黑"
H1_COLOR = "2E74B5"
H2_COLOR = "2E74B5"
H3_COLOR = "1F4D78"
TITLE_COLOR = "0B2545"
MUTED = "595959"
TABLE_HEADER_FILL = "E8EEF5"
TBL_INDENT = 120
CELL_MARGINS = (80, 80, 120, 120)  # top, bottom, start, end


def mean(rows, key):
    vals = [r[key] for r in rows if r.get(key) is not None]
    return round(statistics.mean(vals), 3) if vals else None


def load(path):
    with open(os.path.join(ROOT, path), encoding="utf-8") as f:
        return json.load(f)


# ==================== 数据 ====================
ROUNDS = [
    ("纯模型基线", "纯模型，不检索", "results/eval_results_80_baseline_complete.json"),
    ("纯向量+门控", "纯向量检索（固定切块 800/150）+门控", "results/eval_results_80_adaptive_complete.json"),
    ("leg1（默认）", "混合检索+固定切块 800/150+rerank+门控", "results/eval_results_80_hybrid_oldchunk_complete.json"),
    ("leg2", "混合检索+语义切块+rerank+门控", "results/eval_results_80_hybrid_semantic_complete.json"),
    ("leg3", "混合检索+固定切块，关闭 rerank/门控", "results/eval_results_80_norerank_recursive_complete.json"),
    ("chunk400", "混合检索，切块 400/150", "results/eval_results_80_chunk400_complete.json"),
    ("chunk1200", "混合检索，切块 1200/150", "results/eval_results_80_chunk1200_complete.json"),
    ("回归 200 题", "混合检索，切块 800/150（200 题基准首轮）", "results/regression/eval_results_latest.json"),
]

ROUND_DATA = []
for label, desc, path in ROUNDS:
    d = load(path)
    rows = d["questions"]
    ROUND_DATA.append({
        "label": label,
        "desc": desc,
        "n": len(rows),
        "aq": mean(rows, "answer_quality"),
        "kp": mean(rows, "key_point_coverage"),
        "rs": mean(rows, "retrieval_sufficiency"),
    })

SUBJECTS = ["数据结构", "操作系统", "组成原理", "计算机网络"]
DETAIL_ROUNDS = [
    ("leg1", "results/eval_results_80_hybrid_oldchunk_complete.json"),
    ("chunk400", "results/eval_results_80_chunk400_complete.json"),
    ("chunk1200", "results/eval_results_80_chunk1200_complete.json"),
    ("回归 200 题", "results/regression/eval_results_latest.json"),
]
DETAIL = []
for label, path in DETAIL_ROUNDS:
    d = load(path)
    subj = {}
    for r in d["questions"]:
        subj.setdefault(r["subject"], []).append(r)
    for s in SUBJECTS:
        rows = subj.get(s, [])
        DETAIL.append((label, s, len(rows), mean(rows, "answer_quality"),
                       mean(rows, "key_point_coverage"), mean(rows, "retrieval_sufficiency")))

base80 = load("results/eval_results_80_hybrid_oldchunk_complete.json")["questions"]
reg200 = load("results/regression/eval_results_latest.json")["questions"]
a = {r["id"]: r for r in base80}
b = {r["id"]: r for r in reg200}
common = sorted(set(a) & set(b))
d_aq = [b[i]["answer_quality"] - a[i]["answer_quality"] for i in common
        if a[i].get("answer_quality") is not None and b[i].get("answer_quality") is not None]
d_kp = [b[i]["key_point_coverage"] - a[i]["key_point_coverage"] for i in common
        if a[i].get("key_point_coverage") is not None and b[i].get("key_point_coverage") is not None]
COMMON = {
    "n": len(common),
    "aq": round(statistics.mean(d_aq), 3),
    "aq_up": sum(1 for x in d_aq if x > 0),
    "aq_down": sum(1 for x in d_aq if x < 0),
    "kp": round(statistics.mean(d_kp), 3),
    "kp_up": sum(1 for x in d_kp if x > 0),
    "kp_down": sum(1 for x in d_kp if x < 0),
}


# ==================== 基础工具 ====================
def fmt(v, digits=2):
    return "—" if v is None else f"{v:.{digits}f}"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = BODY_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             before=0, after=6, line=1.25, style=None, indent=None, hang=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent is not None:
        pf.left_indent = Inches(indent)
    if hang is not None:
        pf.first_line_indent = Inches(-hang)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_table_geometry(table, widths):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(TBL_INDENT))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    cellMar = OxmlElement("w:tblCellMar")
    for name, val in zip(("top", "bottom", "start", "end"), CELL_MARGINS):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)
    grid = tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[idx]))
            tcW.set(qn("w:type"), "dxa")
    trPr = table.rows[0]._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def fill_cell(cell, text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(str(text)), size=size, bold=bold, color=color)


def make_table(doc, headers, rows, widths, body_aligns=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for j, h in enumerate(headers):
        fill_cell(table.rows[0].cells[j], h, size=10, bold=True, color=TITLE_COLOR)
        shade_cell(table.rows[0].cells[j], TABLE_HEADER_FILL)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            align = (body_aligns or [WD_ALIGN_PARAGRAPH.CENTER])[j]
            fill_cell(table.rows[i].cells[j], val, size=10, align=align)
    set_table_geometry(table, widths)
    return table


def add_page_field(paragraph, instr):
    r = paragraph.add_run()
    set_run(r, size=9, color=MUTED)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r._r.append(fld1)
    r._r.append(instr_el)
    r._r.append(fld2)


def add_restart_numbering(doc):
    """克隆 List Number 的编号定义，返回新的 numId（供下一段列表从 1 重新计数）。"""
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    abs_nums = numbering.findall(qn("w:abstractNum"))
    src_abs_id = None
    for n in nums:
        if int(n.get(qn("w:numId"))) == 5:  # List Number 样式使用的 numId
            src_abs_id = n.find(qn("w:abstractNumId")).get(qn("w:val"))
            break
    src_abs = next(
        (a for a in abs_nums if a.get(qn("w:abstractNumId")) == src_abs_id), None
    )
    if src_abs is None:
        return None
    new_abs_id = max(int(a.get(qn("w:abstractNumId"))) for a in abs_nums) + 1
    new_num_id = max(int(n.get(qn("w:numId"))) for n in nums) + 1
    clone = copy.deepcopy(src_abs)
    clone.set(qn("w:abstractNumId"), str(new_abs_id))
    nsid = clone.find(qn("w:nsid"))
    if nsid is not None:
        nsid.set(qn("w:val"), "%08X" % (0x40000000 + new_abs_id))
    numbering.append(clone)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(new_abs_id))
    num.append(abs_ref)
    numbering.append(num)
    return new_num_id


def set_paragraph_num(paragraph, num_id, ilvl=0):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    ilvl_el = numPr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = numPr.find(qn("w:numId"))
    if numId_el is None:
        numId_el = OxmlElement("w:numId")
        numPr.append(numId_el)
    numId_el.set(qn("w:val"), str(num_id))


# ==================== 文档 ====================
doc = Document()

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
    setattr(sec, m, Inches(1.0))
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, H1_COLOR, 18, 10),
    ("Heading 2", 13, H2_COLOR, 14, 7),
    ("Heading 3", 12, H3_COLOR, 10, 5),
):
    st = doc.styles[name]
    st.font.name = BODY_FONT
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.25

# 页眉：安静标签 + 细下划线
header_p = sec.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(header_p.add_run("CacheFlow408 · RAG 评测总报告"), size=9, color=MUTED)
pbdr = OxmlElement("w:pBdr")
btm = OxmlElement("w:bottom")
btm.set(qn("w:val"), "single")
btm.set(qn("w:sz"), "4")
btm.set(qn("w:space"), "2")
btm.set(qn("w:color"), "D0D4DA")
pbdr.append(btm)
header_p._p.get_or_add_pPr().append(pbdr)

# 页脚：右侧页码
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer_p.add_run("第 "), size=9, color=MUTED)
add_page_field(footer_p, " PAGE ")
set_run(footer_p.add_run(" 页"), size=9, color=MUTED)

# ==================== 封面标题区 ====================
add_para(doc, "CacheFlow408 评测总报告", size=22, bold=True, color=TITLE_COLOR,
         before=0, after=4)
add_para(doc, "全链路 RAG · 混合检索 · 可评测体系（复习助手 / AI 应用岗简历项目）",
         size=13, color=MUTED, after=8)
add_para(doc, "评测规模：80 题基准 + 200 题回归基准 ｜ 裁判：qwen2.5:7b ｜ 生成：qwen2.5:7b ｜ "
              "检索：bge-m3 + BM25 + RRF + bge-reranker-v2-m3 ｜ 报告日期：2026-08-05",
         size=9.5, color=MUTED, after=4)


# 一、评测体系与方法
doc.add_heading("一、评测体系与方法", level=1)
add_para(doc,
         "本项目围绕“408 考研复习助手”构建了完整的 RAG 链路：查询改写 → 混合检索（向量 + BM25 + RRF 融合）→ "
         "rerank（bge-reranker-v2-m3，fp16）→ 门控降级 → 生成（qwen2.5:7b）。为验证各环节的贡献，建立了可复现的离线评测体系：")
for text in [
    "评测集：80 题（四科各 20 题，含 3 道知识库外超纲题），后扩展为 200 题（四科各 50 题，含 120 道带知识点标签的题库选择题）；",
    "裁判：本地 qwen2.5:7b（temperature=0），一次调用同时给出召回充分性、回答质量与要点命中率；",
    "指标：answer_quality（1-5，结论正确且要点齐备）、key_point_coverage（要点命中率）、retrieval_sufficiency（1-5，检索资料对作答的覆盖充分性）；",
    "工程化：断点续跑、结果归档、多轮对比脚本、200 题一键回归与自动报告。",
]:
    add_para(doc, text, after=3, style="List Bullet", indent=0.375, hang=0.188)


# 二、轮次总览
doc.add_heading("二、轮次总览", level=1)
headers = ["轮次", "配置说明", "题数", "回答质量", "要点命中", "召回充分性"]
rows = [
    (r["label"], r["desc"], r["n"], fmt(r["aq"]), fmt(r["kp"]), fmt(r["rs"]))
    for r in ROUND_DATA
]
widths = [1700, 3260, 700, 1150, 1150, 1400]
make_table(doc, headers, rows, widths,
           body_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
add_para(doc, "表 1：各轮整体指标（数值为 3 位小数均值；— 表示该轮不适用）。",
         size=9, color=MUTED, before=4, after=4)


# 三、关键结论
doc.add_heading("三、关键结论", level=1)
conclusions = [
    "混合检索全面优于纯向量：回答质量 3.79→3.99、要点命中 0.735→0.821、召回充分性 4.25→4.66，四科全胜，作为默认检索方案。",
    "语义切块未胜出：回答质量持平，要点命中 0.797 低于固定切块的 0.821，不作为默认；产品采用固定切块 800/150 + 混合检索。",
    "reranker + 门控保留：关闭 rerank 后回答质量 3.99→3.88、要点命中 0.821→0.784、召回 4.66→4.45；差距主要来自超库题门控（关闭后 3 道超库题全部被迫基于无关资料作答，得 1 分）。",
    "chunk 尺寸 800/150 最优：400（3.96/0.781/4.39）与 1200（3.93/0.768/4.45）均未超越基线，维持默认。",
    "门控误报案例（os-20“活锁”）：reranker 高分 0.687 但内容无关，门控未触发；生成端“上下文不足即拒绝”规则兜底，未产生幻觉但要点得分为 1。单纯调阈值无法干净解决，后续可引入引用溯源与 faithfulness 二次校验。",
    "操作系统题库已补齐答案：426 道由本地模型（带知识点检索接地）判定并写回 md，出题由实时 LLM 判题改为纯题库直出（秒出）。",
]
for text in conclusions:
    add_para(doc, text, after=4, style="List Number", indent=0.375, hang=0.188)


# 四、分科指标明细
doc.add_heading("四、分科指标明细", level=1)
headers2 = ["轮次", "科目", "题数", "回答质量", "要点命中", "召回充分性"]
rows2 = [
    (label, s, n, fmt(aq), fmt(kp), fmt(rs))
    for label, s, n, aq, kp, rs in DETAIL
]
make_table(doc, headers2, rows2, widths,
           body_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
add_para(doc, "表 2：关键轮次分科明细（80 题轮次与 200 题回归基准）。",
         size=9, color=MUTED, before=4, after=4)


# 五、回归与稳定性
doc.add_heading("五、回归与稳定性", level=1)
add_para(doc,
         "2026-08-04 完成 200 题基准首轮回归（默认配置：固定切块 800/150 + 混合检索 + rerank + 门控）。"
         "200 题总体回答质量 3.755、要点命中 0.891、召回充分性 4.18。分科表现：计算机网络最高（回答质量 4.12），"
         "操作系统最低（3.36）——操作系统新增选择题以计算与复杂概念为主，参考答案要求精确选项，得分更严。")
add_para(doc,
         f"与原 80 题基线逐题对比（共同 {COMMON['n']} 题）：回答质量均值差 {COMMON['aq']:+.3f}"
         f"（升 {COMMON['aq_up']}/降 {COMMON['aq_down']}），要点命中均值差 {COMMON['kp']:+.3f}"
         f"（升 {COMMON['kp_up']}/降 {COMMON['kp_down']}），波动处于正常轮次方差范围。"
         "该首轮结果即后续回归的新基线，每次运行 scripts/run_regression.py 自动生成对比报告。")


# 六、门控与引用溯源
doc.add_heading("六、门控与引用溯源", level=1)
add_para(doc,
         "分级降级门控以 rerank 最高分为置信度信号（阈值 0.5）：置信度不足时回退纯模型回答并提示，"
         "避免错误上下文带偏答案。实测超库题 CDN（0.049）与 HTTPS（0.328）均正确触发降级，分别得 5/4 分；"
         "“活锁”题（0.687）为已知误报，由生成端拒绝规则兜底。")
add_para(doc,
         "引用溯源已上线：回答在引用知识点的结论后标注 [资料N]，与参考上下文编号一一对应；"
         "结构化引用包含来源、页码、章节与原文摘要，前端支持点击 [资料N] 定位并高亮对应资料；"
         "模型未标注时后端自动补充“引用依据”行，保证溯源始终可用。")


# 七、结果文件索引
doc.add_heading("七、结果文件索引", level=1)
for text in [
    "纯模型基线：results/eval_results_80_baseline_complete.json",
    "纯向量：results/eval_results_80_adaptive_complete.json",
    "混合检索 + 固定切块（默认 leg1）：results/eval_results_80_hybrid_oldchunk_complete.json",
    "混合检索 + 语义切块：results/eval_results_80_hybrid_semantic_complete.json",
    "关闭 rerank：results/eval_results_80_norerank_recursive_complete.json",
    "chunk 400 / chunk 1200：results/eval_results_80_chunk400_complete.json、_chunk1200_complete.json",
    "200 题回归：results/regression/eval_results_latest.json、results/regression_report.md",
    "评测说明文档：docs/EVAL_REPORT.md",
]:
    add_para(doc, text, after=3, style="List Bullet", indent=0.375, hang=0.188)


# 八、复现方法
doc.add_heading("八、复现方法", level=1)
restart_num_id = add_restart_numbering(doc)
for text in [
    "重建指定切块参数的向量库（可指定库根目录，避免覆盖默认库）：python scripts/rebuild_kb.py --chunk-size 400 --overlap 150 --kb-root storage/chroma_s400",
    "跑完整评测（rerank + 门控，断点续跑）：python evaluate.py --adaptive --resume --questions data/eval_questions_200.json --output results/regression",
    "一键回归并自动生成对比报告：python scripts/run_regression.py",
    "对比多轮结果：python scripts/compare_eval.py results/eval_results_80_hybrid_oldchunk_complete.json results/eval_results_80_chunk400_complete.json",
    "扩展评测基准（80→200 题，固定随机种子）：python scripts/build_benchmark.py",
]:
    p = add_para(doc, text, after=4, style="List Number", indent=0.375, hang=0.188)
    if restart_num_id is not None:
        set_paragraph_num(p, restart_num_id)


os.makedirs(os.path.dirname(os.path.abspath(OUT)) or ".", exist_ok=True)
doc.save(OUT)
print(f"已生成: {OUT}")
